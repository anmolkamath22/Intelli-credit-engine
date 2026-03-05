"""CAM formatting and export helpers."""

from __future__ import annotations

import textwrap
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


def _render_cam_lines(cam_sections: dict) -> list[str]:
    """Render CAM sections into readable lines for TXT/PDF formats."""
    lines: list[str] = []
    lines.append("CREDIT APPRAISAL MEMO (CAM)")
    lines.append("========================================")
    for section, content in cam_sections.items():
        lines.append("")
        lines.append(section.upper())
        lines.append("-" * min(len(section), 40))
        if isinstance(content, list):
            non_empty = [str(x).strip() for x in content if str(x).strip()]
            is_table = bool(non_empty) and all(x.startswith("|") or x.startswith("Financial Performance") for x in non_empty)
            for item in content:
                if is_table:
                    lines.append(str(item))
                    continue
                wrapped = textwrap.wrap(str(item), width=90) or [str(item)]
                first = True
                for w in wrapped:
                    lines.append(f"- {w}" if first else f"  {w}")
                    first = False
        else:
            wrapped = textwrap.wrap(str(content), width=92) or [str(content)]
            lines.extend(wrapped)
    return lines


def _minimal_pdf_from_lines(lines: list[str]) -> bytes:
    """Build a valid multi-page PDF from plain text lines."""
    if not lines:
        lines = ["CREDIT APPRAISAL MEMO (CAM)"]

    max_lines_per_page = 36
    pages: list[list[str]] = []
    page: list[str] = []
    for line in lines:
        for wrapped in textwrap.wrap(line, width=95) or [""]:
            page.append(wrapped)
            if len(page) >= max_lines_per_page:
                pages.append(page)
                page = []
    if page:
        pages.append(page)
    if not pages:
        pages = [["CREDIT APPRAISAL MEMO (CAM)"]]

    def _pdf_escape(s: str) -> str:
        s = s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        return s.encode("latin-1", errors="ignore").decode("latin-1")

    content_streams: list[bytes] = []
    for p in pages:
        content_ops = ["BT", "/F1 10 Tf", "50 760 Td", "14 TL"]
        for i, line in enumerate(p):
            esc = _pdf_escape(line[:200])
            if i == 0:
                content_ops.append(f"({esc}) Tj")
            else:
                content_ops.append(f"T* ({esc}) Tj")
        content_ops.append("ET")
        content_streams.append("\n".join(content_ops).encode("latin-1", errors="ignore"))

    # Object numbering
    # 1: catalog, 2: pages tree, [page/content]*N, last: font
    page_count = len(pages)
    first_page_obj = 3
    font_obj = first_page_obj + page_count * 2

    page_objs: list[bytes] = []
    content_objs: list[bytes] = []
    kids: list[str] = []
    for i in range(page_count):
        page_obj_no = first_page_obj + i * 2
        content_obj_no = page_obj_no + 1
        kids.append(f"{page_obj_no} 0 R")
        page_objs.append(
            (
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                + f"/Contents {content_obj_no} 0 R /Resources << /Font << /F1 {font_obj} 0 R >> >> >>".encode("ascii")
            )
        )
        stream = content_streams[i]
        content_objs.append(
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream"
        )

    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(
        b"<< /Type /Pages /Kids ["
        + " ".join(kids).encode("ascii")
        + b"] /Count "
        + str(page_count).encode("ascii")
        + b" >>"
    )
    for i in range(page_count):
        objects.append(page_objs[i])
        objects.append(content_objs[i])
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray()
    out.extend(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{i} 0 obj\n".encode("ascii"))
        out.extend(obj)
        out.extend(b"\nendobj\n")

    xref_offset = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    out.extend(
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode("ascii")
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode("ascii")
        + b"\n%%EOF\n"
    )
    return bytes(out)


def save_cam_docx(cam_sections: dict, output_path: Path) -> Path:
    """Save CAM sections to DOCX file."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if Document is None:
        fallback = output_path.with_suffix(".txt")
        text = "\n".join(_render_cam_lines(cam_sections))
        fallback.write_text(text, encoding="utf-8")
        return fallback

    doc = Document()
    doc.add_heading("Credit Appraisal Memo (CAM)", level=1)
    for section, content in cam_sections.items():
        doc.add_heading(section, level=2)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(content))
    doc.save(str(output_path))
    return output_path


def save_cam_pdf(cam_content: dict | str, output_path: Path) -> Path:
    """Save CAM summary as formatted PDF file."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(cam_content, dict):
        lines = _render_cam_lines(cam_content)
    else:
        lines = [ln.strip() for ln in str(cam_content).splitlines() if ln.strip()]
    output_path.write_bytes(_minimal_pdf_from_lines(lines))
    return output_path
